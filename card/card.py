from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import random
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

cardNum=20 #设置生成的贺卡数量
nameFile = open("./name.txt", 'r', encoding="UTF-8")
blessingFile = open("./blessing.txt", 'r', encoding="UTF-8")
nameList = nameFile.read()
nameList = nameList.split('、')
blessingList = blessingFile.readlines()
pBlessingList = []
for i in blessingList:
    if i == '\n':
        continue
    pBlessingList.append(re.match(r".*?、(.*)", i).group(1))
# 创建贺卡docx
for i in range(0, cardNum):
    title = "给" + nameList[i] + "的贺卡"
    myDoc = Document()
    myDoc.add_heading(title, 0)
    vis = []
    totalNum = 0
    # 加入四行祝福的话语
    for j in range(0, 50):
        if totalNum == 4:
            break
        pParagraph = random.choice(pBlessingList)
        pVis = pBlessingList.index(pParagraph)
        # 祝福的话语不能重复
        if pVis not in vis:
            vis.append(pBlessingList.index(pParagraph))
            myParagraph = myDoc.add_paragraph(pParagraph)
            myRun = myParagraph.runs[0]
            # 设置首行缩进
            myParagraph.paragraph_format.first_line_indent = Inches(0.3)
            # 设置行间距
            myParagraph.paragraph_format.line_spacing = 1.5
            # 设置字体
            myRun.font.name = u'楷体'
            myRun._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            # 设置字体大小
            myRun.font.size = Pt(16)
            totalNum += 1
        else:
            continue
    # 加入署名
    myParagraph = myDoc.add_paragraph("你的朋友：captainzw")
    # 设置标题署名格式
    # 设置对齐
    myDoc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    myDoc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # 设置字体
    myDoc.paragraphs[0].runs[0].font.name = u'黑体'
    myDoc.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    myDoc.paragraphs[-1].runs[0].font.name = u'黑体'
    myDoc.paragraphs[-1].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    # 设置大小
    myDoc.paragraphs[0].runs[0].font.size = Pt(32)
    myDoc.paragraphs[-1].runs[0].font.size = Pt(18)
    #保存文件
    myDoc.save("./" + title + ".docx")
